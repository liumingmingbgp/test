from docx import Document

doc = Document()
doc.add_heading('一级标题', level=1)
p2 = doc.add_paragraph('第二个段落')
p1 = p2.insert_paragraph_before('第一个段落')
p3 = doc.add_paragraph('新段落')
p3.add_run('加粗').bold = True
p3.add_run('以及')
p3.add_run('斜体').italic = True

doc.save('new1.docx')