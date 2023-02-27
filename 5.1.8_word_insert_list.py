from docx import Document

doc = Document()

style = 'List Number'
doc.add_paragraph('有序列表1', style=style)
doc.add_paragraph('有序列表2', style=style)
doc.add_paragraph('有序列表2', style=style)

style = 'List Bullet'
doc.add_paragraph('无序列表1', style=style)
doc.add_paragraph('无序列表2', style=style)
doc.add_paragraph('无序列表2', style=style)

doc.save('new_list.docx')