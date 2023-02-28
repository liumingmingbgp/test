from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

doc = Document()

P1 = doc.add_paragraph('水平居中对齐')
P1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

P2 = doc.add_paragraph('左对齐')
P2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

P3 = doc.add_paragraph('右对齐')
P3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

P4 = doc.add_paragraph()
run = P4.add_run('内联对象')
run.font.size = Pt(35)
run.font.italic = True

doc.save('new_5.2.1.docx')