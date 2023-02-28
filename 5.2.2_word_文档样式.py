from docx import Document
from docx.shared import Pt
from docx.enum.style import *

doc = Document()
styles = doc.styles

### 使用styles属性定义样式 ###

style = styles['Heading 1']
p_format = style.paragraph_format
p_format.left_indent = Pt(25)  # 设置左缩进磅值
p = doc.add_paragraph('使用style设置段落样式', style=style)
p1 = doc.add_paragraph('hahahahhahahahaha', style=style)
doc.save('use_style.docx')

### 输出pytho-docx库中支持的所有表格样式 ###

for style in styles:
    if style.type == WD_STYLE_TYPE.TABLE:
        doc.add_paragraph(f"表格样式名称：{style.name}")
        table = doc.add_table(3,3, style=style)
        cells = table.rows[0].cells
        cells[0].text = '第一列内容'
        cells[1].text = '第二列内容'
        cells[2].text = '第三列内容'
        doc.add_paragraph('\n')

doc.save('show_all_table_styles.docx')
