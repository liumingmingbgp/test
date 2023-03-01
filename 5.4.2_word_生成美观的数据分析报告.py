from docx import Document
import pandas as pd

student = pd.read_excel('student_score.xlsx')
student.sort_values(by='SCORE', inplace=True, ascending=False)
doc = Document()
doc.add_heading('数据分析报告', level=0)
first_student = student.iloc[0,:]['NAME']
first_score = student.iloc[0,:]['SCORE']

p = doc.add_paragraph('分数排在第一位的同学是：')
p.add_run(str(first_student)).bold=True
p.add_run('，分数为：')
p.add_run(str(first_score)).bold=True

p1 = doc.add_paragraph(f'总共有{len(student["NAME"])}名同学参加了考试，学生考试总体情况：')
table = doc.add_table(rows=len(student['NAME'])+1, cols=2)
table.style = 'LightShading-Accent1'
table.cell(0,0).text = '学生姓名'
table.cell(0,1).text = '学生分数'
for i, (index, row) in enumerate(student.iterrows()):
    table.cell(i+1, 0).text = str(row['NAME'])
    table.cell(i+1, 1).text = str(row['SCORE'])
doc.add_picture('student_score.png')

doc.save('student_scoce_analyze.docx')