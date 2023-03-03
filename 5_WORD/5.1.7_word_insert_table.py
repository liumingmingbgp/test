from docx import Document
from docx.shared import Inches

doc = Document()
table = doc.add_table(rows=3, cols=4)
table.style = 'Table Grid'

row = table.rows[0]
row.cells[0].text = '第一行第一列'

cell = table.cell(0,1)
cell.text = '第一行第二列'

cell1 = table.cell(1,0)
p = cell1.paragraphs[0]
run = p.add_run()
run.add_picture('ICON.jpeg', width=Inches(1.25))

doc.save('new3.docx')